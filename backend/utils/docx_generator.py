from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING


def generate_docx_from_printable(printable_text: str, output_path: str):
    try:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        lines = printable_text.split("\n")

        for line in lines:
            line = line.strip()

            if line.startswith("### "):
                p = doc.add_heading(line.replace("### ", ""), level=2)
                p.paragraph_format.space_after = Pt(4)

            elif line == "":
                p = doc.add_paragraph("")
                p.paragraph_format.space_after = Pt(2)

            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        doc.save(output_path)
        return output_path
    except Exception:
        return ""
